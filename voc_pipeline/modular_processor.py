"""
Modular Pipeline Processor
Allows independent execution of each pipeline stage for maximum flexibility.
"""

import pandas as pd
import os
import sys
import logging
import json
import time
import re
from typing import Dict, List, Optional, Any, Tuple
from datetime import datetime
from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path

# Add the project root to the path
project_root = Path(__file__).parent.parent
sys.path.insert(0, str(project_root))

from langchain.prompts import PromptTemplate
from langchain_openai import ChatOpenAI
from prompts.core_extraction import CORE_EXTRACTION_PROMPT, get_core_extraction_prompt
from prompts.analysis_enrichment import get_analysis_enrichment_prompt

# Set up logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# ===== TIMESTAMP EXTRACTION FUNCTIONS =====
def parse_timestamp(ts_str: str) -> Optional[str]:
    """Parse timestamp string to HH:MM:SS format"""
    if not ts_str:
        return None
    
    # Remove brackets and parentheses
    ts_str = re.sub(r'[\[\]()]', '', ts_str.strip())
    
    # Handle different formats
    patterns = [
        r'^(\d{1,2}):(\d{2}):(\d{2})$',  # HH:MM:SS
        r'^(\d{1,2}):(\d{2})$',          # MM:SS
    ]
    
    for pattern in patterns:
        match = re.match(pattern, ts_str)
        if match:
            if len(match.groups()) == 3:  # HH:MM:SS
                h, m, s = map(int, match.groups())
                return f"{h:02d}:{m:02d}:{s:02d}"
            elif len(match.groups()) == 2:  # MM:SS
                m, s = map(int, match.groups())
                return f"00:{m:02d}:{s:02d}"
    
    return None

def extract_timestamps_from_text(text: str) -> Tuple[Optional[str], Optional[str], List[str]]:
    """
    Extract all timestamps from text and return start, end, and all timestamps found.
    Handles multiple formats: [HH:MM:SS], (MM:SS), Speaker (HH:MM):, Speaker (HH:MM:SS - HH:MM:SS)
    """
    all_timestamps = []
    
    # Pattern 1: [HH:MM:SS] or [MM:SS] inline timestamps
    inline_pattern = r'\[(\d{1,2}:\d{2}(?::\d{2})?)\]'
    for match in re.finditer(inline_pattern, text):
        ts = parse_timestamp(match.group(1))
        if ts:
            all_timestamps.append(ts)
    
    # Pattern 2: Speaker (HH:MM): or (HH:MM):
    speaker_pattern = r'(?:Speaker \d+|[A-Za-z\s]+?)?\s*\((\d{1,2}:\d{2}(?::\d{2})?)\):'
    for match in re.finditer(speaker_pattern, text):
        ts = parse_timestamp(match.group(1))
        if ts:
            all_timestamps.append(ts)
    
    # Pattern 3: ShipBob format - Speaker (HH:MM:SS - HH:MM:SS)
    shipbob_pattern = r'(?:Speaker \d+|[A-Za-z\s]+?)?\s*\((\d{1,2}:\d{2}(?::\d{2})?)\s*-\s*(\d{1,2}:\d{2}(?::\d{2})?)\)'
    for match in re.finditer(shipbob_pattern, text):
        start_ts = parse_timestamp(match.group(1))
        end_ts = parse_timestamp(match.group(2))
        if start_ts:
            all_timestamps.append(start_ts)
        if end_ts:
            all_timestamps.append(end_ts)
    
    # Pattern 4: Standalone (HH:MM): timestamps
    standalone_pattern = r'^\((\d{1,2}:\d{2}(?::\d{2})?)\):'
    for match in re.finditer(standalone_pattern, text, re.MULTILINE):
        ts = parse_timestamp(match.group(1))
        if ts:
            all_timestamps.append(ts)
    
    # Remove duplicates and sort
    all_timestamps = sorted(list(set(all_timestamps)))
    
    start_timestamp = all_timestamps[0] if all_timestamps else None
    end_timestamp = all_timestamps[-1] if all_timestamps else None
    
    return start_timestamp, end_timestamp, all_timestamps

def clean_verbatim_response_with_timestamps(text: str, interviewer_names=None) -> Tuple[str, Optional[str], Optional[str]]:
    """
    Enhanced version of clean_verbatim_response that extracts timestamps before cleaning.
    Returns: (cleaned_text, start_timestamp, end_timestamp)
    """
    # First extract timestamps before cleaning
    start_ts, end_ts, all_ts = extract_timestamps_from_text(text)
    
    # Then clean normally
    if interviewer_names is None:
        interviewer_names = ["Q:", "A:", "Interviewer:", "Drew Giovannoli:", "Brian:", "Yusuf Elmarakby:"]
    
    lines = text.splitlines()
    cleaned_lines = []
    for line in lines:
        l = line.strip()
        # Remove interview titles/headings
        if l.lower().startswith("an interview with") or l.lower().startswith("interview with"):
            continue
        # Remove speaker labels, timestamps, Q:/A: tags
        if any(l.startswith(name) for name in interviewer_names):
            continue
        if re.match(r'^Speaker \d+ \(\d{1,2}:\d{2}(?::\d{2})?\):', l):
            continue
        if re.match(r'^\(\d{1,2}:\d{2}(?::\d{2})?\):', l):
            continue
        # Skip lines that are just questions (but keep questions that are part of longer content)
        if l.endswith("?") and len(l) < 50:
            continue
        cleaned_lines.append(line)
    
    cleaned = " ".join(cleaned_lines).strip()
    
    # Remove timestamps and speaker labels from cleaned text
    cleaned = re.sub(r'\[(\d{1,2}:\d{2}(?::\d{2})?)\]', '', cleaned)  # Remove [HH:MM:SS]
    cleaned = re.sub(r'Speaker \d+ \(\d{1,2}:\d{2}(?::\d{2})?\):\s*', '', cleaned)  # Remove Speaker (HH:MM):
    cleaned = re.sub(r'\(\d{1,2}:\d{2}(?::\d{2})?\):\s*', '', cleaned)  # Remove (HH:MM):
    cleaned = re.sub(r'(?:Speaker \d+|[A-Za-z\s]+?)?\s*\(\d{1,2}:\d{2}(?::\d{2})?\s*-\s*\d{1,2}:\d{2}(?::\d{2})?\)', '', cleaned)  # Remove ShipBob format
    cleaned = re.sub(r'^(Speaker \d+|Drew Giovannoli|Brian|Yusuf Elmarakby):\s*', '', cleaned, flags=re.MULTILINE)
    
    # Clean up whitespace
    cleaned = re.sub(r'\n+', ' ', cleaned)
    cleaned = re.sub(r'\s+', ' ', cleaned).strip()
    
    if len(cleaned) < 5:
        cleaned = ""
    
    return cleaned, start_ts, end_ts
# ===== END TIMESTAMP EXTRACTION FUNCTIONS =====



class ModularProcessor:
    """Modular processor for independent pipeline stages."""
    
    def __init__(self, model_name: str = "gpt-4o-mini", max_tokens: int = 4096, temperature: float = 0.3):
        self.model_name = model_name
        self.max_tokens = max_tokens
        self.temperature = temperature
        
        # Initialize LLM
        self.llm = ChatOpenAI(
            model_name=model_name,
            max_tokens=max_tokens,
            temperature=temperature
        )
        
        # Initialize database if available
        self.db = None
        # The original code had DB_AVAILABLE, but DB_AVAILABLE was removed from imports.
        # Assuming DB_AVAILABLE is no longer needed or will be re-added.
        # For now, commenting out the DB initialization as DB_AVAILABLE is gone.
        # if DB_AVAILABLE:
        #     try:
        #         from database import VOCDatabase
        #         self.db = VOCDatabase()
        #     except Exception as e:
        #         logger.warning(f"Database not available: {e}")
    
    def stage1_core_extraction(self, transcript_path: str, company: str, interviewee: str, 
                              deal_status: str, date_of_interview: str) -> List[Dict]:
        """
        Stage 1: Core extraction - extract verbatim responses and metadata only.
        
        Returns:
            List of dictionaries with core fields only
        """
        logger.info(f"Starting Stage 1: Core extraction from {transcript_path}")
        
        # Load transcript
        if transcript_path.lower().endswith(".docx"):
            try:
                # Try python-docx first for better extraction
                from docx import Document
                doc = Document(transcript_path)
                full_text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                if not full_text.strip():
                    # Fallback to Docx2txtLoader
                    from langchain_community.document_loaders import Docx2txtLoader
                    loader = Docx2txtLoader(transcript_path)
                    docs = loader.load()
                    full_text = docs[0].page_content
            except ImportError:
                # Fallback to Docx2txtLoader if python-docx not available
                from langchain_community.document_loaders import Docx2txtLoader
                loader = Docx2txtLoader(transcript_path)
                docs = loader.load()
                full_text = docs[0].page_content
        else:
            with open(transcript_path, encoding="utf-8") as f:
                full_text = f.read()
        
        if not full_text.strip():
            raise ValueError("Transcript is empty")
        
        # Debug: Check text length
        import tiktoken
        encoding = tiktoken.get_encoding("cl100k_base")
        total_tokens = len(encoding.encode(full_text))
        logger.info(f"Extracted {len(full_text)} characters ({total_tokens} tokens) from transcript")
        logger.info(f"Text preview: {full_text[:200]}...")
        
        # Create chunks using smaller chunks for better reliability and coverage
        chunks = self._create_chunks(full_text, target_tokens=1000, overlap_tokens=200)
        logger.info(f"Passing {len(chunks)} chunks to LLM with smaller chunks targeting ~2-3 insights per chunk")
        
        # Use parallel processing for better performance
        return self._process_chunks_parallel(chunks, company, interviewee, deal_status, date_of_interview)

    def stage1_core_extraction_sequential(self, transcript_path: str, company: str, interviewee: str, 
                                        deal_status: str, date_of_interview: str) -> List[Dict]:
        """
        Stage 1: Core extraction - SEQUENTIAL VERSION (fallback).
        This preserves the original sequential implementation.
        
        Returns:
            List of dictionaries with core fields only
        """
        logger.info(f"Starting Stage 1: Core extraction (SEQUENTIAL) from {transcript_path}")
        
        # Load transcript
        if transcript_path.lower().endswith(".docx"):
            try:
                # Try python-docx first for better extraction
                from docx import Document
                doc = Document(transcript_path)
                full_text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                if not full_text.strip():
                    # Fallback to Docx2txtLoader
                    from langchain_community.document_loaders import Docx2txtLoader
                    loader = Docx2txtLoader(transcript_path)
                    docs = loader.load()
                    full_text = docs[0].page_content
            except ImportError:
                # Fallback to Docx2txtLoader if python-docx not available
                from langchain_community.document_loaders import Docx2txtLoader
                loader = Docx2txtLoader(transcript_path)
                docs = loader.load()
                full_text = docs[0].page_content
        else:
            with open(transcript_path, encoding="utf-8") as f:
                full_text = f.read()
        
        if not full_text.strip():
            raise ValueError("Transcript is empty")
        
        # Debug: Check text length
        import tiktoken
        encoding = tiktoken.get_encoding("cl100k_base")
        total_tokens = len(encoding.encode(full_text))
        logger.info(f"Extracted {len(full_text)} characters ({total_tokens} tokens) from transcript")
        logger.info(f"Text preview: {full_text[:200]}...")
        
        # Create chunks using smaller chunks for better reliability and coverage
        chunks = self._create_chunks(full_text, target_tokens=1000, overlap_tokens=200)
        logger.info(f"Passing {len(chunks)} chunks to LLM with smaller chunks targeting ~2-3 insights per chunk")
        
        # Process each chunk SEQUENTIALLY (original implementation)
        all_responses = []
        for i, chunk in enumerate(chunks):
            try:
                # Create unique response ID for this chunk
                chunk_id = f"{company}_{interviewee}_{i+1}"
                
                # Create prompt template (use raw template, not formatted)
                prompt_template = PromptTemplate(
                    input_variables=["response_id", "company", "interviewee_name", "deal_status", "date_of_interview", "chunk_text", "start_timestamp", "end_timestamp"],
                    template=CORE_EXTRACTION_PROMPT
                )
                chain = prompt_template | self.llm
                
                # Get response
                result = chain.invoke({
                    "response_id": chunk_id,
                    "company": company,
                    "interviewee_name": interviewee,
                    "deal_status": deal_status,
                    "date_of_interview": date_of_interview,
                    "chunk_text": chunk
                })
                
                # Parse response
                if hasattr(result, 'content'):
                    response_text = result.content.strip()
                else:
                    response_text = str(result).strip()
                
                # Debug: Print raw LLM output
                logger.info(f"LLM raw output for chunk {i}: {repr(response_text[:500])}...")
                
                # Parse the response
                parsed_responses = self._parse_llm_response(response_text, chunk_id, i)
                all_responses.extend(parsed_responses)
                
            except Exception as e:
                logger.error(f"Error processing chunk {i}: {e}")
                continue
        
        # Post-processing: Remove duplicates and improve quality
        return self._post_process_responses(all_responses)

    def _process_chunks_parallel(self, chunks: List[str], company: str, interviewee: str, 
                               deal_status: str, date_of_interview: str, max_workers: int = 3) -> List[Dict]:
        """
        Process chunks in parallel for faster Stage 1 extraction.
        
        Args:
            chunks: List of text chunks to process
            company: Company name
            interviewee: Interviewee name
            deal_status: Deal status
            date_of_interview: Date of interview
            max_workers: Maximum number of parallel workers (default: 3)
            
        Returns:
            List of processed responses
        """
        logger.info(f"🚀 Processing {len(chunks)} chunks in parallel with {max_workers} workers")
        start_time = time.time()
        
        all_responses = []
        
        def process_single_chunk(chunk_info):
            """Process a single chunk - designed to be thread-safe"""
            chunk_index, chunk_text = chunk_info
            try:
                # Create unique response ID for this chunk
                chunk_id = f"{company}_{interviewee}_{chunk_index+1}"
                
                # Extract timestamps from chunk before processing
                start_ts, end_ts, all_ts = extract_timestamps_from_text(chunk_text)
                
                # Clean the chunk text with timestamp extraction
                cleaned_chunk, clean_start_ts, clean_end_ts = clean_verbatim_response_with_timestamps(chunk_text)
                
                # Use cleaned text for LLM processing
                processed_chunk = cleaned_chunk if cleaned_chunk else chunk_text
                
                # Create prompt template (use raw template, not formatted)
                prompt_template = PromptTemplate(
                    input_variables=["response_id", "company", "interviewee_name", "deal_status", "date_of_interview", "chunk_text", "start_timestamp", "end_timestamp"],
                    template=CORE_EXTRACTION_PROMPT
                )
                chain = prompt_template | self.llm
                
                # Get response with timestamps
                result = chain.invoke({
                    "response_id": chunk_id,
                    "company": company,
                    "interviewee_name": interviewee,
                    "deal_status": deal_status,
                    "date_of_interview": date_of_interview,
                    "chunk_text": processed_chunk,
                    "start_timestamp": clean_start_ts or start_ts,
                    "end_timestamp": clean_end_ts or end_ts
                })
                
                # Parse response
                if hasattr(result, 'content'):
                    response_text = result.content.strip()
                else:
                    response_text = str(result).strip()
                
                # Debug: Print raw LLM output
                logger.info(f"🔍 LLM raw output for chunk {chunk_index}: {repr(response_text[:200])}...")
                
                # Parse the response
                parsed_responses = self._parse_llm_response(response_text, chunk_id, chunk_index)
                
                # Add timestamps to all responses (LLM might not include them consistently)
                for response in parsed_responses:
                    if 'start_timestamp' not in response or not response.get('start_timestamp'):
                        response["start_timestamp"] = clean_start_ts or start_ts or "00:00:00"
                    if 'end_timestamp' not in response or not response.get('end_timestamp'):
                        response["end_timestamp"] = clean_end_ts or end_ts or "00:00:00"
                
                logger.info(f"✅ Chunk {chunk_index} completed: {len(parsed_responses)} responses extracted")
                return parsed_responses
                
            except Exception as e:
                logger.error(f"❌ Error processing chunk {chunk_index}: {e}")
                return []
        
        # Prepare chunk information for parallel processing
        chunk_info_list = [(i, chunk) for i, chunk in enumerate(chunks)]
        
        # Process chunks in parallel
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            # Submit all chunk processing tasks
            future_to_chunk = {
                executor.submit(process_single_chunk, chunk_info): chunk_info[0] 
                for chunk_info in chunk_info_list
            }
            
            # Collect results as they complete
            for future in as_completed(future_to_chunk):
                chunk_index = future_to_chunk[future]
                try:
                    chunk_responses = future.result()
                    all_responses.extend(chunk_responses)
                except Exception as exc:
                    logger.error(f"❌ Chunk {chunk_index} generated an exception: {exc}")
        
        processing_time = time.time() - start_time
        logger.info(f"🎉 Parallel processing completed in {processing_time:.2f} seconds")
        logger.info(f"📊 Extracted {len(all_responses)} total responses from {len(chunks)} chunks")
        
        # Post-processing: Remove duplicates and improve quality
        return self._post_process_responses(all_responses)

    def _parse_llm_response(self, response_text: str, chunk_id: str, chunk_index: int) -> List[Dict]:
        """
        Parse LLM response text into structured data.
        Extracted from the original sequential implementation for reuse.
        """
        parsed_responses = []
        
        # Fix JSON parsing: Strip markdown code blocks if present
        if response_text.startswith('```json'):
            response_text = response_text[7:]  # Remove ```json
        if response_text.startswith('```'):
            response_text = response_text[3:]  # Remove ```
        if response_text.endswith('```'):
            response_text = response_text[:-3]  # Remove trailing ```
        response_text = response_text.strip()
        
        # Parse JSON
        try:
            responses = json.loads(response_text)
            if isinstance(responses, list):
                # Add chunk index to response IDs to ensure uniqueness
                for j, response in enumerate(responses):
                    if 'response_id' in response:
                        response['response_id'] = f"{chunk_id}_{j+1}"
                    # Output validation for 'question' field
                    question = response.get('question', '')
                    if not self._is_valid_question(question):
                        logger.warning(f"[Stage1 Extraction] Invalid or missing question for response_id {response.get('response_id')}: '{question}'. Setting to 'UNKNOWN'.")
                        response['question'] = 'UNKNOWN'
                parsed_responses.extend(responses)
            else:
                responses['response_id'] = f"{chunk_id}_1"
                # Output validation for 'question' field
                question = responses.get('question', '')
                if not self._is_valid_question(question):
                    logger.warning(f"[Stage1 Extraction] Invalid or missing question for response_id {responses.get('response_id')}: '{question}'. Setting to 'UNKNOWN'.")
                    responses['question'] = 'UNKNOWN'
                parsed_responses.append(responses)
        except json.JSONDecodeError as e:
            logger.warning(f"Failed to parse JSON from chunk {chunk_index}: {e}")
            logger.warning(f"Raw response was: {repr(response_text)}")
        
        return parsed_responses

    def _post_process_responses(self, all_responses: List[Dict]) -> List[Dict]:
        """
        Post-process responses to remove duplicates and improve quality.
        Extracted from the original implementation for reuse.
        """
        logger.info(f"📋 Post-processing {len(all_responses)} responses")
        
        # Remove exact duplicates based on verbatim response
        seen_responses = set()
        unique_responses = []
        for response in all_responses:
            verbatim = response.get('verbatim_response', '').strip().lower()
            if verbatim and verbatim not in seen_responses:
                seen_responses.add(verbatim)
                unique_responses.append(response)
        
        logger.info(f"🔄 Removed {len(all_responses) - len(unique_responses)} duplicate responses")
        
        # Filter out very short responses (likely incomplete) - reduced threshold for better coverage
        quality_responses = []
        for response in unique_responses:
            verbatim = response.get('verbatim_response', '')
            if len(verbatim.split()) >= 10:  # Reduced from 20 to 10 words for better coverage
                quality_responses.append(response)
        
        logger.info(f"✂️ Post-processing removed {len(unique_responses) - len(quality_responses)} short responses")
        logger.info(f"✅ Stage 1 complete: extracted {len(quality_responses)} quality responses")
        
        return quality_responses
    
    def stage2_analysis_enrichment(self, stage1_data_responses: List[Dict]) -> List[Dict]:
        """
        Stage 2: Analysis enrichment - add AI-generated insights to core responses.
        
        Args:
            stage1_data_responses: List of dictionaries with core fields
            
        Returns:
            List of dictionaries with core fields + analysis fields
        """
        logger.info(f"Starting Stage 2: Analysis enrichment for {len(stage1_data_responses)} responses")
        
        enriched_responses = []
        
        for response in stage1_data_responses:
            try:
                # Create analysis prompt
                prompt_text = get_analysis_enrichment_prompt(
                    response_id=response['response_id'],
                    verbatim_response=response['verbatim_response'],
                    subject=response['subject'],
                    company=response['company'],
                    interviewee_name=response['interviewee_name']
                )
                
                # Create chain
                prompt_template = PromptTemplate(
                    input_variables=["response_id", "verbatim_response", "subject", "company", "interviewee_name"],
                    template=prompt_text
                )
                chain = prompt_template | self.llm
                
                # Get analysis
                result = chain.invoke({
                    "response_id": response['response_id'],
                    "verbatim_response": response['verbatim_response'],
                    "subject": response['subject'],
                    "company": response['company'],
                    "interviewee_name": response['interviewee_name']
                })
                
                # Parse response
                if hasattr(result, 'content'):
                    analysis_text = result.content.strip()
                else:
                    analysis_text = str(result).strip()
                
                # Parse JSON
                try:
                    analysis = json.loads(analysis_text)
                    
                    # Combine core response with analysis
                    enriched_response = {**response, **analysis}
                    enriched_responses.append(enriched_response)
                    
                except json.JSONDecodeError as e:
                    logger.warning(f"Failed to parse analysis JSON for {response['response_id']}: {e}")
                    # Keep original response without analysis
                    enriched_responses.append(response)
                
            except Exception as e:
                logger.error(f"Error enriching response {response['response_id']}: {e}")
                # Keep original response without analysis
                enriched_responses.append(response)
        
        logger.info(f"Stage 2 complete: enriched {len(enriched_responses)} responses")
        return enriched_responses
    
    def stage3_labeling(self, responses: List[Dict]) -> List[Dict]:
        """
        Stage 3: Labeling - add structured labels to responses.
        
        Args:
            responses: List of dictionaries with core fields (and optionally analysis fields)
            
        Returns:
            List of dictionaries with core fields + labels
        """
        logger.info(f"Starting Stage 3: Labeling for {len(responses)} responses")
        
        labeled_responses = []
        
        for response in responses:
            try:
                # Create labeling prompt
                prompt_text = get_labeling_prompt(
                    response_id=response['response_id'],
                    verbatim_response=response['verbatim_response'],
                    subject=response['subject'],
                    company=response['company'],
                    interviewee_name=response['interviewee_name']
                )
                
                # Create chain
                prompt_template = PromptTemplate(
                    input_variables=["response_id", "verbatim_response", "subject", "company", "interviewee_name"],
                    template=prompt_text
                )
                chain = prompt_template | self.llm
                
                # Get labels
                result = chain.invoke({
                    "response_id": response['response_id'],
                    "verbatim_response": response['verbatim_response'],
                    "subject": response['subject'],
                    "company": response['company'],
                    "interviewee_name": response['interviewee_name']
                })
                
                # Parse response
                if hasattr(result, 'content'):
                    labels_text = result.content.strip()
                else:
                    labels_text = str(result).strip()
                
                # Parse JSON
                try:
                    labels_data = json.loads(labels_text)
                    
                    # Combine response with labels
                    labeled_response = {**response, 'labels': labels_data.get('labels', {})}
                    labeled_responses.append(labeled_response)
                    
                except json.JSONDecodeError as e:
                    logger.warning(f"Failed to parse labels JSON for {response['response_id']}: {e}")
                    # Keep original response without labels
                    labeled_responses.append(response)
                
            except Exception as e:
                logger.error(f"Error labeling response {response['response_id']}: {e}")
                # Keep original response without labels
                labeled_responses.append(response)
        
        logger.info(f"Stage 3 complete: labeled {len(labeled_responses)} responses")
        return labeled_responses
    
    def save_to_database(self, responses: List[Dict]) -> int:
        """
        Save responses to database.
        
        Args:
            responses: List of response dictionaries
            
        Returns:
            Number of responses saved
        """
        if not self.db:
            logger.warning("Database not available")
            return 0
        
        saved_count = 0
        for response in responses:
            try:
                if self.db.save_response(response):
                    saved_count += 1
                else:
                    logger.warning(f"Failed to save response {response.get('response_id', 'unknown')}")
            except Exception as e:
                logger.error(f"Error saving response {response.get('response_id', 'unknown')}: {e}")
        
        logger.info(f"Saved {saved_count} responses to database")
        return saved_count
    
    def _create_chunks(self, text: str, target_tokens: int = 1000, overlap_tokens: int = 400) -> List[str]:
        """
        Create text chunks for processing using enhanced Q&A-aware chunking with better overlap.
        
        Args:
            text: Full transcript text
            target_tokens: Target tokens per chunk (1K for better reliability)
            overlap_tokens: Overlap between chunks (400 tokens for better Q&A preservation)
            
        Returns:
            List of text chunks optimized for comprehensive processing with enhanced Q&A preservation
        """
        import tiktoken
        
        # Initialize tokenizer
        try:
            encoding = tiktoken.encoding_for_model("gpt-4o-mini")
        except:
            # Fallback to cl100k_base encoding
            encoding = tiktoken.get_encoding("cl100k_base")
        
        # Tokenize the full text
        tokens = encoding.encode(text)
        
        chunks = []
        start = 0
        
        while start < len(tokens):
            end = start + target_tokens
            
            # Extract tokens for this chunk
            chunk_tokens = tokens[start:end]
            
            # Decode back to text
            chunk_text = encoding.decode(chunk_tokens)
            
            # Enhanced Q&A-aware segmentation - try to break at natural conversation boundaries
            if end < len(tokens):
                # Look for natural break points in the remaining text with extended lookahead
                remaining_tokens = tokens[end:end + 500]  # Increased look ahead to 500 tokens
                remaining_text = encoding.decode(remaining_tokens)
                
                # Find the best break point
                break_point = self._find_break_point(chunk_text, remaining_text)
                if break_point > len(chunk_text) * 0.6:  # More flexible break point (60% instead of 70%)
                    chunk_text = chunk_text[:break_point]
                    # Adjust end position based on actual text length
                    actual_tokens = encoding.encode(chunk_text)
                    end = start + len(actual_tokens)
            
            chunks.append(chunk_text)
            
            # Move start position with enhanced overlap
            start = end - overlap_tokens
            if start >= len(tokens):
                break
        
        logger.info(f"Created {len(chunks)} chunks with enhanced Q&A-aware token-based chunking")
        logger.info(f"Average chunk size: {sum(len(encoding.encode(chunk)) for chunk in chunks) // len(chunks)} tokens")
        logger.info(f"Enhanced overlap: {overlap_tokens} tokens for better Q&A preservation")
        
        return chunks
    
    def _find_break_point(self, chunk_text: str, remaining_text: str) -> int:
        """
        Find the best break point for enhanced Q&A-aware segmentation.
        
        Args:
            chunk_text: Current chunk text
            remaining_text: Next portion of text
            
        Returns:
            Best break point position in chunk_text
        """
        # Enhanced conversation boundaries with more patterns
        break_patterns = [
            '\n\n',  # Double newline (speaker change)
            '.\n',   # End of sentence followed by newline
            '?\n',   # Question followed by newline
            '!\n',   # Exclamation followed by newline
            '\nQ:',  # Question marker
            '\nA:',  # Answer marker
            '\nInterviewer:',  # Interviewer marker
            '\nInterviewee:',  # Interviewee marker
            '\nDrew Giovannoli:',  # Specific interviewer name
            '\nCyrus Nazarian:',   # Specific interviewee name
            '\nModerator:',        # Generic moderator
            '\nSpeaker:',          # Generic speaker
            '.\n\n',              # End of sentence with double newline
            '?\n\n',              # Question with double newline
            '!\n\n',              # Exclamation with double newline
        ]
        
        best_break = len(chunk_text)
        
        # Look for break points in the latter 40% of the chunk (more flexible than 70%)
        search_start = int(len(chunk_text) * 0.6)
        
        for pattern in break_patterns:
            pos = chunk_text.rfind(pattern, search_start)
            if pos > search_start:
                best_break = min(best_break, pos + len(pattern))
        
        # If no good break point found, try to break at sentence boundaries
        if best_break == len(chunk_text):
            # Look for sentence endings in the latter part
            sentence_endings = ['. ', '? ', '! ', '.\n', '?\n', '!\n']
            for ending in sentence_endings:
                pos = chunk_text.rfind(ending, search_start)
                if pos > search_start:
                    best_break = min(best_break, pos + len(ending))
        
        return best_break
    
    def _is_valid_question(self, question: str) -> bool:
        """Check if a string is a valid question (interrogative) with enhanced complex question support."""
        if not question or not isinstance(question, str):
            return False
        q = question.strip()
        
        # Must end with a question mark or start with a question word
        question_words = ("what", "how", "why", "when", "who", "where", "which", "is", "are", "do", "does", "did", "can", "could", "would", "should", "will", "has", "have", "had")
        
        # Basic validation
        if q.endswith("?") or q.lower().startswith(question_words):
            return True
        
        # Enhanced validation for complex questions
        # Check for questions embedded in statements
        q_lower = q.lower()
        
        # Look for question patterns within complex statements
        complex_patterns = [
            "were there any",
            "did you",
            "do you",
            "would you",
            "could you",
            "should you",
            "are there",
            "is there",
            "have you",
            "has there",
            "will you",
            "can you",
            "might you",
            "must you",
            "need you",
            "want you",
            "expect you",
            "plan to",
            "consider",
            "think about",
            "evaluate",
            "compare",
            "assess",
            "review",
            "analyze"
        ]
        
        for pattern in complex_patterns:
            if pattern in q_lower:
                return True
        
        # Check for questions that start with statements but end with questions
        # Example: "Speed and cost were priorities. Were there other criteria?"
        if "." in q and any(word in q.split(".")[-1].lower() for word in question_words):
            return True
        
        return False
    
    def run_full_pipeline(self, transcript_path: str, company: str, interviewee: str, 
                         deal_status: str, date_of_interview: str, 
                         save_to_db: bool = True) -> Dict[str, Any]:
        """
        Run the complete pipeline (all stages).
        
        Returns:
            Dictionary with results from each stage
        """
        logger.info("Starting full pipeline")
        
        # Stage 1: Core extraction
        stage1_data_responses = self.stage1_core_extraction(
            transcript_path, company, interviewee, deal_status, date_of_interview
        )
        
        # Stage 2: Analysis enrichment
        enriched_responses = self.stage2_analysis_enrichment(stage1_data_responses)
        
        # Stage 3: Labeling
        labeled_responses = self.stage3_labeling(enriched_responses)
        
        # Save to database if requested
        saved_count = 0
        if save_to_db:
            saved_count = self.save_to_database(labeled_responses)
        
        return {
            'stage1_core_count': len(stage1_data_responses),
            'stage2_enriched_count': len(enriched_responses),
            'stage3_labeled_count': len(labeled_responses),
            'database_saved_count': saved_count,
            'responses': labeled_responses
        } 